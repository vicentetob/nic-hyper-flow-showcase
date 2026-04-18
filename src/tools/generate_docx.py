#!/usr/bin/env python3
"""
generate_docx.py
Gera arquivos .docx a partir de um JSON spec recebido via stdin ou --spec.

Uso:
  echo '{"output_path": "doc.docx", ...}' | python generate_docx.py
  python generate_docx.py --spec '{"output_path": "doc.docx", ...}'
  python generate_docx.py --file spec.json

Retorno (stdout, JSON):
  {"ok": true, "path": "/abs/path/to/doc.docx"}
  {"ok": false, "error": "mensagem de erro"}
"""

import sys
import json
import argparse
import traceback
from pathlib import Path

# ── dependências ────────────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import docx.opc.constants
except ImportError:
    print(json.dumps({
        "ok": False, 
        "error": "Biblioteca 'python-docx' não encontrada no ambiente.",
        "suggestion": "Instale a dependência necessária executando: pip install python-docx",
        "missing_dependency": "python-docx"
    }))
    sys.exit(1)

# ── helpers ──────────────────────────────────────────────────────────────────

ALIGN_MAP = {
    "left":    WD_ALIGN_PARAGRAPH.LEFT,
    "center":  WD_ALIGN_PARAGRAPH.CENTER,
    "right":   WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

HEADING_LEVELS = {"h1": 1, "h2": 2, "h3": 3, "h4": 4, "h5": 5, "h6": 6}

def hex_to_rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

def apply_run_fmt(run, fmt: dict):
    """Aplica formatação em um run: bold, italic, underline, size, color, font, highlight."""
    if fmt.get("bold"):       run.bold = True
    if fmt.get("italic"):     run.italic = True
    if fmt.get("underline"):  run.underline = True
    if fmt.get("strike"):     run.font.strike = True
    if fmt.get("size"):       run.font.size = Pt(fmt["size"])
    if fmt.get("color"):      run.font.color.rgb = hex_to_rgb(fmt["color"])
    if fmt.get("font"):       run.font.name = fmt["font"]
    if fmt.get("highlight"):
        # highlight aceita nomes de cor WD_COLOR_INDEX
        from docx.enum.text import WD_COLOR_INDEX
        color_map = {
            "yellow": WD_COLOR_INDEX.YELLOW, "green": WD_COLOR_INDEX.GREEN,
            "cyan": WD_COLOR_INDEX.CYAN, "magenta": WD_COLOR_INDEX.MAGENTA,
            "blue": WD_COLOR_INDEX.BLUE, "red": WD_COLOR_INDEX.RED,
            "dark_blue": WD_COLOR_INDEX.DARK_BLUE, "teal": WD_COLOR_INDEX.TEAL,
            "dark_yellow": WD_COLOR_INDEX.DARK_YELLOW,
        }
        run.font.highlight_color = color_map.get(fmt["highlight"].lower())

def set_paragraph_fmt(para, fmt: dict):
    """Aplica formatação de parágrafo: align, spacing, indent."""
    if fmt.get("align"):
        para.alignment = ALIGN_MAP.get(fmt["align"].lower(), WD_ALIGN_PARAGRAPH.LEFT)
    pf = para.paragraph_format
    if fmt.get("space_before"):  pf.space_before = Pt(fmt["space_before"])
    if fmt.get("space_after"):   pf.space_after  = Pt(fmt["space_after"])
    if fmt.get("line_spacing"):  pf.line_spacing  = Pt(fmt["line_spacing"])
    if fmt.get("left_indent"):   pf.left_indent   = Inches(fmt["left_indent"])
    if fmt.get("first_line"):    pf.first_line_indent = Inches(fmt["first_line"])

def add_hyperlink(para, url: str, text: str, fmt: dict = None):
    """Insere um hyperlink clicável num parágrafo existente."""
    part = para.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    para._p.append(hyperlink)
    if fmt:
        # aplica formatação adicional ao run dentro do hyperlink
        from docx.text.run import Run
        run_obj = Run(new_run, para)
        apply_run_fmt(run_obj, fmt)

def add_runs(para, runs_spec: list):
    """
    runs_spec pode ser:
      - list of strings: ["texto simples"]
      - list of dicts:   [{"text": "...", "bold": true, "color": "#FF0000"}]
      - list misto
      - dict com "url" para hyperlink: {"url": "https://...", "text": "clique aqui"}
    """
    for item in runs_spec:
        if isinstance(item, str):
            para.add_run(item)
        elif isinstance(item, dict):
            if item.get("url"):
                add_hyperlink(para, item["url"], item.get("text", item["url"]), item)
            elif item.get("page_break"):
                run = para.add_run()
                run.add_break(docx.enum.text.WD_BREAK.PAGE)
            else:
                run = para.add_run(item.get("text", ""))
                apply_run_fmt(run, item)

# ── builders de blocos ────────────────────────────────────────────────────────

def build_paragraph(doc: Document, block: dict):
    """
    block = {
      "type": "paragraph",
      "text": "string simples",       # OU
      "runs": [...],                   # runs formatados
      "style": "Normal",              # estilo word (opcional)
      "fmt": { align, size, ... }     # formatação extra
    }
    """
    style = block.get("style", "Normal")
    try:
        para = doc.add_paragraph(style=style)
    except Exception:
        para = doc.add_paragraph()

    if block.get("text"):
        run = para.add_run(block["text"])
        if block.get("fmt"):
            apply_run_fmt(run, block.get("fmt", {}))
    elif block.get("runs"):
        add_runs(para, block["runs"])

    if block.get("fmt"):
        set_paragraph_fmt(para, block["fmt"])
    return para

def build_heading(doc: Document, block: dict):
    """
    block = {
      "type": "h1" | "h2" | ... | "h6",
      "text": "Título",
      "runs": [...],
      "fmt": {...}
    }
    """
    level = HEADING_LEVELS.get(block["type"], 1)
    para = doc.add_heading(level=level)
    para.clear()  # remove conteúdo padrão

    if block.get("text"):
        para.add_run(block["text"])
    elif block.get("runs"):
        add_runs(para, block["runs"])

    if block.get("fmt"):
        set_paragraph_fmt(para, block["fmt"])
    return para

def build_list(doc: Document, block: dict):
    """
    block = {
      "type": "list",
      "ordered": false,
      "items": [
        "texto simples",
        {"text": "item formatado", "bold": true},
        {"runs": [...]},
        {"text": "sub-item", "level": 1}
      ]
    }
    """
    ordered = block.get("ordered", False)
    style = "List Number" if ordered else "List Bullet"
    for item in block.get("items", []):
        level = 0
        if isinstance(item, dict):
            level = item.get("level", 0)

        # sub-níveis têm estilos diferentes
        if level == 0:
            s = style
        else:
            s = ("List Number %d" % (level + 1)) if ordered else ("List Bullet %d" % (level + 1))

        try:
            para = doc.add_paragraph(style=s)
        except Exception:
            para = doc.add_paragraph(style=style)

        if isinstance(item, str):
            para.add_run(item)
        elif isinstance(item, dict):
            if item.get("runs"):
                add_runs(para, item["runs"])
            else:
                run = para.add_run(item.get("text", ""))
                apply_run_fmt(run, item)

def build_table(doc: Document, block: dict):
    """
    block = {
      "type": "table",
      "headers": ["Col A", "Col B"],     # opcional (linha de header)
      "rows": [
        ["Célula 1", "Célula 2"],
        [{"text": "Bold", "bold": true}, "Normal"]
      ],
      "style": "Table Grid",             # estilo Word (opcional)
      "col_widths": [2.5, 4.0],          # largura em inches (opcional)
      "header_bg": "#2E75B6",            # cor de fundo do header (opcional)
      "header_color": "#FFFFFF"          # cor do texto do header (opcional)
    }
    """
    headers = block.get("headers", [])
    rows_data = block.get("rows", [])
    num_cols = len(headers) or (len(rows_data[0]) if rows_data else 1)

    total_rows = len(rows_data) + (1 if headers else 0)
    style = block.get("style", "Table Grid")
    try:
        table = doc.add_table(rows=total_rows, cols=num_cols, style=style)
    except Exception:
        table = doc.add_table(rows=total_rows, cols=num_cols)

    # larguras de coluna
    col_widths = block.get("col_widths", [])
    if col_widths:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                if i < len(col_widths):
                    cell.width = Inches(col_widths[i])

    row_offset = 0

    # header
    if headers:
        hrow = table.rows[0]
        hrow.height = Inches(0.35)
        for i, htext in enumerate(headers):
            cell = hrow.cells[i]
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(htext)
            run.bold = True
            if block.get("header_color"):
                run.font.color.rgb = hex_to_rgb(block["header_color"])
            if block.get("header_bg"):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), block["header_bg"].lstrip("#"))
                tcPr.append(shd)
        row_offset = 1

    # dados
    for ri, row_data in enumerate(rows_data):
        trow = table.rows[row_offset + ri]
        for ci, cell_data in enumerate(row_data):
            if ci >= num_cols:
                break
            cell = trow.cells[ci]
            cell.text = ""
            para = cell.paragraphs[0]
            if isinstance(cell_data, str):
                para.add_run(cell_data)
            elif isinstance(cell_data, dict):
                if cell_data.get("runs"):
                    add_runs(para, cell_data["runs"])
                else:
                    run = para.add_run(cell_data.get("text", ""))
                    apply_run_fmt(run, cell_data)
                if cell_data.get("align"):
                    para.alignment = ALIGN_MAP.get(cell_data["align"], WD_ALIGN_PARAGRAPH.LEFT)
                if cell_data.get("bg"):
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"), "clear")
                    shd.set(qn("w:color"), "auto")
                    shd.set(qn("w:fill"), cell_data["bg"].lstrip("#"))
                    tcPr.append(shd)

    return table

def build_image(doc: Document, block: dict):
    """
    block = {
      "type": "image",
      "path": "/caminho/para/imagem.png",
      "width": 4.0,     # inches (opcional)
      "height": 3.0,    # inches (opcional)
      "align": "center" # opcional
    }
    """
    path = block.get("path", "")
    if not Path(path).exists():
        para = doc.add_paragraph(f"[Imagem não encontrada: {path}]")
        return

    kwargs = {}
    if block.get("width"):  kwargs["width"]  = Inches(block["width"])
    if block.get("height"): kwargs["height"] = Inches(block["height"])

    para = doc.add_paragraph()
    if block.get("align"):
        para.alignment = ALIGN_MAP.get(block["align"].lower(), WD_ALIGN_PARAGRAPH.LEFT)
    run = para.add_run()
    run.add_picture(path, **kwargs)

def build_page_break(doc: Document, block: dict):
    doc.add_page_break()

def build_hr(doc: Document, block: dict):
    """Linha horizontal via borda inferior de parágrafo vazio."""
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(block.get("thickness", 6)))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), block.get("color", "000000").lstrip("#"))
    pBdr.append(bottom)
    pPr.append(pBdr)

def build_toc_placeholder(doc: Document, block: dict):
    """Insere um parágrafo de Sumário (TOC) — Word atualiza ao abrir."""
    para = doc.add_paragraph()
    run = para.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar)
    run2 = para.add_run()
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._r.append(instrText)
    run3 = para.add_run()
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    run3._r.append(fldChar2)
    run4 = para.add_run()
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run4._r.append(fldChar3)

# ── dispatcher ────────────────────────────────────────────────────────────────

BUILDERS = {
    "paragraph":   build_paragraph,
    "p":           build_paragraph,
    "h1": build_heading, "h2": build_heading, "h3": build_heading,
    "h4": build_heading, "h5": build_heading, "h6": build_heading,
    "list":        build_list,
    "table":       build_table,
    "image":       build_image,
    "img":         build_image,
    "page_break":  build_page_break,
    "hr":          build_hr,
    "toc":         build_toc_placeholder,
}

def build_block(doc: Document, block: dict):
    btype = block.get("type", "paragraph").lower()
    builder = BUILDERS.get(btype)
    if builder:
        builder(doc, block)
    else:
        # fallback: trata como parágrafo
        build_paragraph(doc, {**block, "type": "paragraph"})

# ── page setup ────────────────────────────────────────────────────────────────

def apply_page_setup(doc: Document, page: dict):
    """
    page = {
      "size": "A4" | "letter",
      "landscape": false,
      "margin_top": 1.0,     # inches
      "margin_bottom": 1.0,
      "margin_left": 1.25,
      "margin_right": 1.25
    }
    """
    from docx.shared import Inches
    section = doc.sections[0]

    size = page.get("size", "A4").upper()
    landscape = page.get("landscape", False)

    if size == "LETTER":
        w, h = Inches(8.5), Inches(11)
    elif size == "A3":
        w, h = Cm(29.7), Cm(42)
    else:  # A4
        w, h = Cm(21), Cm(29.7)

    if landscape:
        section.page_width, section.page_height = h, w
        section.orientation = 1  # LANDSCAPE
    else:
        section.page_width, section.page_height = w, h
        section.orientation = 0  # PORTRAIT

    if page.get("margin_top")    is not None: section.top_margin    = Inches(page["margin_top"])
    if page.get("margin_bottom") is not None: section.bottom_margin = Inches(page["margin_bottom"])
    if page.get("margin_left")   is not None: section.left_margin   = Inches(page["margin_left"])
    if page.get("margin_right")  is not None: section.right_margin  = Inches(page["margin_right"])

# ── header / footer ───────────────────────────────────────────────────────────

def apply_header_footer(doc: Document, spec: dict, kind: str):
    """
    spec = {
      "text": "texto simples",
      "runs": [...],
      "align": "center",
      "page_number": true   # insere número de página
    }
    kind = "header" | "footer"
    """
    section = doc.sections[0]
    obj = section.header if kind == "header" else section.footer
    obj.is_linked_to_previous = False

    para = obj.paragraphs[0] if obj.paragraphs else obj.add_paragraph()
    para.clear()
    if spec.get("align"):
        para.alignment = ALIGN_MAP.get(spec["align"].lower(), WD_ALIGN_PARAGRAPH.LEFT)

    if spec.get("text"):
        para.add_run(spec["text"])
    elif spec.get("runs"):
        add_runs(para, spec["runs"])

    if spec.get("page_number"):
        run = para.add_run()
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        run._r.append(fldChar1)
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = " PAGE "
        run._r.append(instrText)
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar2)

# ── default styles ────────────────────────────────────────────────────────────

def apply_default_styles(doc: Document, styles: dict):
    """
    styles = {
      "font": "Calibri",
      "size": 11,
      "heading1": {"size": 18, "color": "#1F3864", "bold": true},
      "heading2": {"size": 14, "color": "#2E75B6", "bold": true}
    }
    """
    if styles.get("font") or styles.get("size"):
        style = doc.styles["Normal"]
        if styles.get("font"):  style.font.name = styles["font"]
        if styles.get("size"):  style.font.size = Pt(styles["size"])

    for level in range(1, 7):
        key = f"heading{level}"
        if styles.get(key):
            hs = styles[key]
            try:
                hstyle = doc.styles[f"Heading {level}"]
                if hs.get("size"):  hstyle.font.size = Pt(hs["size"])
                if hs.get("bold") is not None: hstyle.font.bold = hs["bold"]
                if hs.get("color"): hstyle.font.color.rgb = hex_to_rgb(hs["color"])
                if hs.get("font"):  hstyle.font.name = hs["font"]
            except Exception:
                pass

# ── entry point ───────────────────────────────────────────────────────────────

def generate(spec: dict) -> dict:
    output_path = spec.get("output_path", "output.docx")
    output_path = str(Path(output_path).expanduser().resolve())

    append_to = spec.get("append_to")
    appending = False

    if append_to:
        # ── modo append: abre o docx existente e anexa nova página ──────────
        src = str(Path(append_to).expanduser().resolve())
        if not Path(src).exists():
            return {"ok": False, "error": f"append_to não encontrado: {src}"}
        doc = Document(src)
        appending = True

        # quebra de página antes do novo conteúdo
        doc.add_page_break()

        # header/footer opcionais por seção — se o chamador quiser override
        if spec.get("header"):
            apply_header_footer(doc, spec["header"], "header")
        if spec.get("footer"):
            apply_header_footer(doc, spec["footer"], "footer")

    else:
        # ── modo criação: documento novo ─────────────────────────────────────
        doc = Document()

        # page setup
        if spec.get("page"):
            apply_page_setup(doc, spec["page"])
        else:
            apply_page_setup(doc, {"size": "A4"})

        # estilos globais
        if spec.get("styles"):
            apply_default_styles(doc, spec["styles"])

        # header / footer
        if spec.get("header"):
            apply_header_footer(doc, spec["header"], "header")
        if spec.get("footer"):
            apply_header_footer(doc, spec["footer"], "footer")

        # remover parágrafo vazio inicial do python-docx
        if doc.paragraphs:
            p = doc.paragraphs[0]._element
            p.getparent().remove(p)

    # ── blocos de conteúdo (comum aos dois modos) ────────────────────────────
    for block in spec.get("content", []):
        try:
            build_block(doc, block)
        except Exception as e:
            doc.add_paragraph(f"[ERRO no bloco {block.get('type','?')}: {e}]")

    doc.save(output_path)
    return {
        "ok":      True,
        "path":    output_path,
        "appended": appending,
        "pages_so_far": spec.get("page_index", "?"),
    }


def main():
    parser = argparse.ArgumentParser(description="Gerador de .docx via JSON spec")
    parser.add_argument("--spec", help="JSON spec como string")
    parser.add_argument("--file", help="Caminho para arquivo JSON spec")
    args = parser.parse_args()

    if args.spec:
        raw = args.spec
    elif args.file:
        raw = Path(args.file).read_text(encoding="utf-8")
    else:
        raw = sys.stdin.read()

    try:
        spec = json.loads(raw)
    except json.JSONDecodeError as e:
        print(json.dumps({"ok": False, "error": f"JSON inválido: {e}"}))
        sys.exit(1)

    try:
        result = generate(spec)
        print(json.dumps(result))
    except Exception as e:
        print(json.dumps({"ok": False, "error": traceback.format_exc()}))
        sys.exit(1)


if __name__ == "__main__":
    main()
